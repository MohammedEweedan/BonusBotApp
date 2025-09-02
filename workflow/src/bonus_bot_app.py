import os
import platform
import subprocess
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
import threading
import sys

class BonusBotApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Bonus Bot - Document Generator")
        self.root.geometry("600x500")
        self.root.resizable(True, True)
        
        # File paths
        self.excel_path = ""
        self.template_path = ""
        self.output_folder = ""
        
        self.setup_ui()
        
    def setup_ui(self):
        """Setup the user interface"""
        # Main frame
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        # Title
        title_label = ttk.Label(main_frame, text="Bonus Bot Document Generator", 
                               font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        # Excel file selection
        ttk.Label(main_frame, text="Excel File (.xlsx):").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.excel_label = ttk.Label(main_frame, text="No file selected", 
                                    foreground="gray")
        self.excel_label.grid(row=1, column=1, sticky=(tk.W, tk.E), padx=(10, 10))
        ttk.Button(main_frame, text="Browse", 
                  command=self.browse_excel).grid(row=1, column=2, padx=(5, 0))
        
        # Template file selection
        ttk.Label(main_frame, text="Template (.docx):").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.template_label = ttk.Label(main_frame, text="No file selected", 
                                       foreground="gray")
        self.template_label.grid(row=2, column=1, sticky=(tk.W, tk.E), padx=(10, 10))
        ttk.Button(main_frame, text="Browse", 
                  command=self.browse_template).grid(row=2, column=2, padx=(5, 0))
        
        # Output folder selection
        ttk.Label(main_frame, text="Output Folder:").grid(row=3, column=0, sticky=tk.W, pady=5)
        self.output_label = ttk.Label(main_frame, text="No folder selected", 
                                     foreground="gray")
        self.output_label.grid(row=3, column=1, sticky=(tk.W, tk.E), padx=(10, 10))
        ttk.Button(main_frame, text="Browse", 
                  command=self.browse_output).grid(row=3, column=2, padx=(5, 0))
        
        # Generate button
        self.generate_btn = ttk.Button(main_frame, text="Generate Documents", 
                                      command=self.start_generation, state="disabled")
        self.generate_btn.grid(row=4, column=0, columnspan=3, pady=20)
        
        # Progress bar
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progress.grid(row=5, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
        
        # Status text
        self.status_text = tk.Text(main_frame, height=15, width=70)
        self.status_text.grid(row=6, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(10, 0))
        
        # Scrollbar for status text
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=self.status_text.yview)
        scrollbar.grid(row=6, column=3, sticky=(tk.N, tk.S))
        self.status_text.configure(yscrollcommand=scrollbar.set)
        
        # Configure grid weights for resizing
        main_frame.rowconfigure(6, weight=1)
        
    def browse_excel(self):
        """Browse for Excel file"""
        file_path = filedialog.askopenfilename(
            title="Select Excel File",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")]
        )
        if file_path:
            self.excel_path = file_path
            self.excel_label.config(text=os.path.basename(file_path), foreground="black")
            self.check_ready()
            
    def browse_template(self):
        """Browse for template file"""
        file_path = filedialog.askopenfilename(
            title="Select Template File",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
        )
        if file_path:
            self.template_path = file_path
            self.template_label.config(text=os.path.basename(file_path), foreground="black")
            self.check_ready()
            
    def browse_output(self):
        """Browse for output folder"""
        folder_path = filedialog.askdirectory(title="Select Output Folder")
        if folder_path:
            self.output_folder = folder_path
            self.output_label.config(text=folder_path, foreground="black")
            self.check_ready()
            
    def check_ready(self):
        """Check if all files are selected and enable generate button"""
        if self.excel_path and self.template_path and self.output_folder:
            self.generate_btn.config(state="normal")
        else:
            self.generate_btn.config(state="disabled")
            
    def log_message(self, message):
        """Add message to status text"""
        self.status_text.insert(tk.END, message + "\n")
        self.status_text.see(tk.END)
        self.root.update_idletasks()
        
    def start_generation(self):
        """Start the document generation in a separate thread"""
        self.generate_btn.config(state="disabled")
        self.progress.start()
        self.status_text.delete(1.0, tk.END)
        
        # Run generation in separate thread to prevent UI freezing
        thread = threading.Thread(target=self.generate_documents)
        thread.daemon = True
        thread.start()
        
    def replace_placeholders_bold(self, doc, replacements):
        """Replace placeholders in document while preserving formatting"""
        replacements_made = []
        
        for para in doc.paragraphs:
            # Get all text from all runs in the paragraph
            full_text = "".join(run.text for run in para.runs)
            
            # Check if any placeholder exists in this paragraph
            replacement_made = False
            modified_text = full_text
            
            # Replace all placeholders in the text
            for key, val in replacements.items():
                if key in modified_text:
                    modified_text = modified_text.replace(key, str(val))
                    replacement_made = True
                    replacements_made.append(f"Replaced {key} with {val}")
            
            # If replacements were made, update the paragraph
            if replacement_made:
                # Clear existing runs
                para.clear()
                # Add the modified text as a new run
                run = para.add_run(modified_text)
                run.bold = True
                
        return replacements_made
        
    def generate_doc(self, employee, output_path):
        """Generate a Word document for an employee using the template"""
        doc = Document(self.template_path)

        # Get today's date in DD/MM/YYYY format
        today = datetime.today().strftime("%d/%m/%Y")

        # Prepare replacements dictionary
        replacements = {
            "{{current_date}}": today,
            "{{full_name}}": str(employee["full_name"]) if pd.notna(employee["full_name"]) else "",
            "{{first_name}}": str(employee["first_name"]) if pd.notna(employee["first_name"]) else "",
            "{{job_description}}": str(employee["job_description"]) if pd.notna(employee["job_description"]) else "",
            "{{branch}}": str(employee["branch"]) if pd.notna(employee["branch"]) else "",
            "{{branch_grams_0000}}": str(int(employee["branch_grams_0000"])) if pd.notna(employee["branch_grams_0000"]) else "",
            "{{personal_grams_0000}}": str(int(employee["personal_grams_0000"])) if pd.notna(employee["personal_grams_0000"]) else "",
            "{{dinar}}": str(int(employee["dinar"])) if pd.notna(employee["dinar"]) else "",
            "{{value_18ct}}": str(int(employee["value_18ct"])) if pd.notna(employee["value_18ct"]) else "",
            "{{value_21ct}}": str(int(employee["value_21ct"])) if pd.notna(employee["value_21ct"]) else "",
        }
        
        # Apply replacements to main document
        replacements_made = self.replace_placeholders_bold(doc, replacements)
        
        # Also check headers and footers for placeholders
        for section in doc.sections:
            # Check header
            if section.header:
                for para in section.header.paragraphs:
                    full_text = "".join(run.text for run in para.runs)
                    modified_text = full_text
                    replacement_made = False
                    
                    for key, val in replacements.items():
                        if key in modified_text:
                            modified_text = modified_text.replace(key, str(val))
                            replacement_made = True
                            replacements_made.append(f"Replaced {key} with {val} in header")
                    
                    if replacement_made:
                        para.clear()
                        run = para.add_run(modified_text)
                        run.bold = True
            
            # Check footer
            if section.footer:
                for para in section.footer.paragraphs:
                    full_text = "".join(run.text for run in para.runs)
                    modified_text = full_text
                    replacement_made = False
                    
                    for key, val in replacements.items():
                        if key in modified_text:
                            modified_text = modified_text.replace(key, str(val))
                            replacement_made = True
                            replacements_made.append(f"Replaced {key} with {val} in footer")
                    
                    if replacement_made:
                        para.clear()
                        run = para.add_run(modified_text)
                        run.bold = True

        # Adjust page margins and font size
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        # Reduce font size to conserve space
        for para in doc.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)

        doc.save(output_path)
        return replacements_made
        
    def convert_to_pdf(self, docx_path, pdf_path):
        """Convert Word document to PDF"""
        system = platform.system()

        if system == "Windows":
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False

                doc = word.Documents.Open(os.path.abspath(docx_path))
                doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
                doc.Close()
                word.Quit()
                return True
            except Exception as e:
                self.log_message(f"❌ Error with Word conversion: {str(e)}")
                # Try LibreOffice as fallback
                return self.convert_with_libreoffice(docx_path, pdf_path)
        else:
            return self.convert_with_libreoffice(docx_path, pdf_path)
            
    def convert_with_libreoffice(self, docx_path, pdf_path):
        """Convert using LibreOffice as fallback"""
        try:
            # Try common LibreOffice paths
            libreoffice_paths = [
                "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
                "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                "libreoffice",
                "soffice"
            ]
            
            libreoffice_path = None
            for path in libreoffice_paths:
                if os.path.exists(path) or path in ["libreoffice", "soffice"]:
                    libreoffice_path = path
                    break
                    
            if not libreoffice_path:
                self.log_message("❌ LibreOffice not found. Please install LibreOffice or Microsoft Word.")
                return False

            subprocess.run([
                libreoffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", os.path.dirname(pdf_path),
                docx_path
            ], check=True, capture_output=True)
            return True
            
        except Exception as e:
            self.log_message(f"❌ Error with LibreOffice conversion: {str(e)}")
            return False
        
    def generate_documents(self):
        """Main document generation function"""
        try:
            # Read the Excel file
            df = pd.read_excel(self.excel_path)
            self.log_message(f"📊 Found {len(df)} employees in Excel file")
            
            # Check if required columns exist
            required_columns = ['full_name', 'dinar']
            missing_columns = [col for col in required_columns if col not in df.columns]
            if missing_columns:
                self.log_message(f"❌ Missing required columns: {missing_columns}")
                self.log_message(f"Available columns: {list(df.columns)}")
                messagebox.showerror("Error", f"Missing required columns: {missing_columns}")
                return

            # Create output subfolder with current month/year
            now = datetime.today()
            month = now.strftime("%B")
            year = now.strftime("%Y")
            month_year = f"{month} {year}"
            folder_name = f"{month_year} - Bonus Letters"
            output_path = os.path.join(self.output_folder, folder_name)
            os.makedirs(output_path, exist_ok=True)
            
            self.log_message(f"📁 Created output folder: {folder_name}")
            
            success_count = 0
            error_count = 0
            
            for index, row in df.iterrows():
                employee = row.to_dict()
                
                # Skip rows with missing full_name
                if pd.isna(employee.get('full_name')):
                    self.log_message(f"⚠️ Skipping row {index + 1}: Missing full_name")
                    continue
                    
                # Filename with underscores and month-year suffix
                base_name = f"{employee['full_name'].replace(' ', '_')}_{month}_{year}"
                docx_path = os.path.join(output_path, f"{base_name}.docx")
                pdf_path = os.path.join(output_path, f"{base_name}.pdf")

                try:
                    self.log_message(f"📝 Processing: {employee['full_name']}")
                    
                    # Generate Word document
                    replacements_made = self.generate_doc(employee, docx_path)
                    
                    # Convert to PDF
                    if self.convert_to_pdf(docx_path, pdf_path):
                        os.remove(docx_path)  # Clean up the temporary Word file
                        self.log_message(f"✅ Generated: {base_name}.pdf")
                        success_count += 1
                    else:
                        self.log_message(f"⚠️ Word document created but PDF conversion failed for {employee['full_name']}")
                        error_count += 1
                        
                except Exception as e:
                    self.log_message(f"❌ Error processing {employee['full_name']}: {str(e)}")
                    error_count += 1

            # Final summary
            self.log_message(f"\n🎉 Generation Complete!")
            self.log_message(f"✅ Successfully generated: {success_count} documents")
            if error_count > 0:
                self.log_message(f"❌ Errors: {error_count} documents")
            self.log_message(f"📁 Output location: {output_path}")
            
            # Show completion message
            messagebox.showinfo("Complete", 
                              f"Document generation complete!\n"
                              f"Success: {success_count}\n"
                              f"Errors: {error_count}\n\n"
                              f"Files saved to:\n{output_path}")

        except FileNotFoundError:
            error_msg = f"Excel file not found: {self.excel_path}"
            self.log_message(f"❌ {error_msg}")
            messagebox.showerror("Error", error_msg)
        except Exception as e:
            error_msg = f"Error reading Excel file: {str(e)}"
            self.log_message(f"❌ {error_msg}")
            messagebox.showerror("Error", error_msg)
        finally:
            self.progress.stop()
            self.generate_btn.config(state="normal")


def main():
    """Main function to run the application"""
    root = tk.Tk()
    app = BonusBotApp(root)
    
    # Center the window on screen
    root.update_idletasks()
    x = (root.winfo_screenwidth() // 2) - (root.winfo_width() // 2)
    y = (root.winfo_screenheight() // 2) - (root.winfo_height() // 2)
    root.geometry(f"+{x}+{y}")
    
    root.mainloop()


if __name__ == "__main__":
    main()